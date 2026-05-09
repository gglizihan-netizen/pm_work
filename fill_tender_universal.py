#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件通用填充脚本 - 基于通用指南
自适应不同格式，无需硬编码索引
"""

from docx import Document
from docx.oxml.ns import qn
import re

# ============================================
# 第一部分：通用工具库
# ============================================

def get_paragraph_text(para):
    """获取段落的完整文本"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def set_paragraph_text(para, new_text):
    """设置段落的完整文本"""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)

def get_cell_text(cell):
    """获取单元格文本"""
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""

def set_cell_text(cell, text):
    """安全设置单元格文本"""
    if cell.paragraphs:
        first_para = cell.paragraphs[0]
        if first_para.runs:
            first_para.runs[0].text = text
            for run in first_para.runs[1:]:
                run.text = ''
        else:
            first_para.text = text
    else:
        cell.text = text

def get_grid_span(cell):
    """获取单元格跨越的列数"""
    tc = cell._tc
    tcPr = tc.tcPr
    if tcPr is not None:
        gridSpan_elem = tcPr.gridSpan
        if gridSpan_elem is not None:
            return int(gridSpan_elem.val)
    return 1

def find_tables_by_header(doc, header_keywords):
    """
    通过表头关键词查找表格
    支持两种模式：1) 表头行包含关键词  2) 表格内任意单元格包含关键词
    """
    results = []
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # 方法1：检查前两行作为表头
        for row_idx in range(min(2, len(table.rows))):
            row_text = " ".join([get_cell_text(cell) for cell in table.rows[row_idx].cells])
            for kw in header_keywords:
                if kw in row_text:
                    results.append((i, kw))
                    break
            if results and results[-1][0] == i:
                break

        # 方法2：如果表头没找到，检查整个表格内容（用于没有明显表头的表格）
        if not any(r[0] == i for r in results):
            all_text = ""
            for row in table.rows[:3]:  # 检查前3行
                all_text += " ".join([get_cell_text(cell) for cell in row.cells])

            for kw in header_keywords:
                if kw in all_text:
                    results.append((i, kw))
                    break

    return results

# ============================================
# 第二部分：配置映射（可扩展）
# ============================================

BRACKET_PLACEHOLDERS = {
    '招标人名称': ['（招标人名称）', '（甲方）', '（采购人）', '（建设单位）'],
    '招标项目名称': ['（招标项目名称）', '（项目名称）', '（工程名称）'],
    '投标人名称': ['（投标人名称）', '（供应商名称）', '（乙方）', '（承包单位）'],
    '制造商名称': ['（制造商名称）', '（生产厂家）', '（厂商名称）'],
    '制造商地址': ['（制造商地址）', '（厂家地址）'],
    '标段编号': ['（标段编号）', '（项目编号）', '（招标编号）'],
    '设备名称': ['（设备名称）', '（产品名称）', '（货物名称）'],
}

LABEL_PATTERHOLDERS = {
    '投标人名称': [['投 标 人：', '供 应 商：', '乙    方：'], '投标人名称'],
    '单位性质': [['单位性质：', '企业性质：'], '单位性质'],
    '地址': [['地    址：', '单位地址：', '注册地址：'], '地址'],
    '成立时间': [['成立时间：', '设立日期：'], '成立时间'],
    '经营期限': [['经营期限：', '营业期限：'], '经营期限'],
    '邮政编码': [['邮政编码：', '邮    编：'], '邮政编码'],
}

TABLE_HEADER_KEYWORDS = {
    '投标人基本情况表': ['投标人名称', '注册资金', '成立时间', '注册地址'],  # 通过内容标签识别
    '法定代表人信息表': ['姓', '名', '性别', '年龄', '职务'],  # 法定代表人表格的特征标签
    '项目情况表': ['设备名称', '规格型号', '项目业主', '合同签订时间'],
    '项目负责人简历': ['姓名', '年龄', '职务', '职称'],
}

TABLE_FIELD_LABELS = {
    '投标人名称': ['投标人名称', '供应商名称', '单位名称'],
    '注册资金': ['注册资金', '注册资本'],
    '成立时间': ['成立时间', '设立日期'],
    '地址': ['注册地址', '单位地址', '地址'],
    '邮政编码': ['邮政编码', '邮编'],
    '员工总数': ['员工总数', '职工人数'],
    '法定代表人姓名': ['姓名'],
    '法定代表人年龄': ['年龄'],
    '项目负责人姓名': ['姓名'],
    '项目负责人年龄': ['年龄'],
}

# ============================================
# 第三部分：通用处理函数
# ============================================

def fill_bracket_placeholders(doc, data):
    """填充括号型占位符"""
    fill_count = {}

    # 遍历段落
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        for data_key, placeholder_list in BRACKET_PLACEHOLDERS.items():
            if data_key not in data:
                continue
            for placeholder in placeholder_list:
                if placeholder in text:
                    new_text = text.replace(placeholder, data[data_key])
                    set_paragraph_text(para, new_text)
                    fill_count[data_key] = fill_count.get(data_key, 0) + 1
                    break

    # 遍历表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = get_cell_text(cell)
                for data_key, placeholder_list in BRACKET_PLACEHOLDERS.items():
                    if data_key not in data:
                        continue
                    for placeholder in placeholder_list:
                        if placeholder in text:
                            new_text = text.replace(placeholder, data[data_key])
                            set_cell_text(cell, new_text)
                            fill_count[data_key] = fill_count.get(data_key, 0) + 1
                            break

    return fill_count

def fill_date_placeholders(doc, year, month, day):
    """填充日期占位符"""
    replacements = {
        '_年_月_日': f"{year}年{month}月{day}日",
        '_年 _月 _日': f"{year}年 {month}月 {day}日",
        '_年': f"{year}年",
        '_月': f"{month}月",
        '_日': f"{day}日",
    }

    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(para, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = get_cell_text(cell)
                new_text = text
                for old, new in replacements.items():
                    new_text = new_text.replace(old, new)
                if new_text != text:
                    set_cell_text(cell, new_text)

def fill_label_placeholders(doc, data):
    """填充标签型占位符"""
    for para in doc.paragraphs:
        text = get_paragraph_text(para).strip()
        for data_key, (label_list, data_source) in LABEL_PATTERHOLDERS.items():
            for label in label_list:
                if text == label or (text.startswith(label) and len(text) <= len(label) + 2):
                    new_value = label + data.get(data_source, '')
                    set_paragraph_text(para, new_value)
                    break

def fill_table_smart(table, data):
    """智能表格填充"""
    fill_results = {}

    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell_text = get_cell_text(cell).strip()

            # 检查是否是标签单元格
            for data_key, label_list in TABLE_FIELD_LABELS.items():
                if data_key not in data:
                    continue

                is_match = any(label in cell_text for label in label_list)

                if is_match:
                    # 找到标签后的下一个单元格
                    next_idx = ci + 1
                    while next_idx < len(row.cells):
                        target_cell = row.cells[next_idx]
                        target_text = get_cell_text(target_cell).strip()

                        # 如果该单元格是另一个标签，跳过
                        is_another_label = any(
                            l in target_text
                            for labels in TABLE_FIELD_LABELS.values()
                            for l in labels
                        )

                        if not is_another_label:
                            set_cell_text(target_cell, data[data_key])
                            fill_results[data_key] = True
                            break

                        next_idx += 1

    return fill_results

def fill_specific_table(doc, table_keywords, data):
    """填充特定表格"""
    tables = find_tables_by_header(doc, table_keywords)

    if not tables:
        return False

    for table_idx, matched_kw in tables:
        table = doc.tables[table_idx]
        fill_table_smart(table, data)

    return True

# ============================================
# 第四部分：文档分析工具
# ============================================

def analyze_document(doc_path):
    """分析文档结构"""
    doc = Document(doc_path)

    print("=" * 80)
    print("文档结构分析报告")
    print("=" * 80)
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")

    # 分析括号型占位符
    bracket_pattern = re.compile(r'（[^）]+）')
    placeholders = set()
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        matches = bracket_pattern.findall(text)
        placeholders.update(matches)

    print(f"\n发现的括号型占位符 ({len(placeholders)}个):")
    for p in sorted(placeholders):
        print(f"  - {p}")

    # 分析表格
    print(f"\n表格结构:")
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        print(f"\n  表格 {i}: {len(table.rows)}行")
        for ri in range(min(3, len(table.rows))):
            row_content = []
            for ci, cell in enumerate(table.rows[ri].cells[:5]):  # 只显示前5列
                text = get_cell_text(cell)[:15]
                span = get_grid_span(cell)
                span_info = f"s{span}" if span > 1 else ""
                row_content.append(f"[{ci}]{text}{span_info}")
            print(f"    行{ri}: {' | '.join(row_content)}")

# ============================================
# 第五部分：主处理流程
# ============================================

def process_document(input_path, output_path, data):
    """主处理函数"""
    print(f"\n正在处理: {input_path}")
    doc = Document(input_path)

    # 步骤1: 填充括号型占位符
    print("\n步骤1: 填充括号型占位符...")
    results = fill_bracket_placeholders(doc, data)
    print(f"  替换了 {sum(results.values())} 处，涉及 {len(results)} 种占位符")

    # 步骤2: 填充日期
    print("\n步骤2: 填充日期...")
    fill_date_placeholders(
        doc,
        data.get('投标日期年', '2025'),
        data.get('投标日期月', '01'),
        data.get('投标日期日', '15')
    )

    # 步骤3: 填充标签型占位符
    print("\n步骤3: 填充标签型占位符...")
    fill_label_placeholders(doc, data)

    # 步骤4: 填充表格
    print("\n步骤4: 智能填充表格...")
    for table_name, keywords in TABLE_HEADER_KEYWORDS.items():
        found = fill_specific_table(doc, keywords, data)
        status = "找到并处理" if found else "未找到"
        print(f"  {table_name}: {status}")

    # 保存
    doc.save(output_path)
    print(f"\n已保存到: {output_path}")

# ============================================
# 第六部分：示例数据
# ============================================

SAMPLE_DATA = {
    '招标人名称': '华美建设集团有限公司',
    '招标项目名称': '华美大厦电梯采购及安装项目',
    '投标人名称': '深圳市华升电梯设备有限公司',
    '制造商名称': '上海三菱电梯有限公司',
    '制造商地址': '上海市闵行区紫竹科学园区江川路811号',
    '标段编号': 'T-2025-001',
    '设备名称': '乘客电梯 HLK-1000/1.75',

    '单位性质': '有限责任公司',
    '地址': '深圳市南山区科技园南区高新南一道001号华升大厦18层',
    '成立时间': '2010年08月18日',
    '经营期限': '2010年08月18日至2030年08月17日',
    '注册资金': '人民币捌仟万元整',
    '邮政编码': '518057',
    '员工总数': '268人',

    '法定代表人姓名': '张伟',
    '法定代表人年龄': '48',
    '项目负责人姓名': '王建国',
    '项目负责人年龄': '45',

    '投标日期年': '2025',
    '投标日期月': '01',
    '投标日期日': '15',
}

if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == '--analyze':
        # 分析模式
        doc_path = sys.argv[2] if len(sys.argv) > 2 else r"D:\桌面\招标文件正文.docx"
        analyze_document(doc_path)
    else:
        # 填充模式
        input_file = r"D:\桌面\招标文件正文.docx"
        output_file = r"D:\桌面\招标文件正文_已填充_通用版.docx"
        process_document(input_file, output_file, SAMPLE_DATA)
