#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件自动填充脚本 - 基于python-docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# ============ 虚拟企业投标数据 ============
COMPANY_DATA = {
    # 核心企业信息
    '投标人名称': '深圳市华升电梯设备有限公司',
    '招标人名称': '华美建设集团有限公司',
    '招标项目名称': '华美大厦电梯采购及安装项目',
    '制造商名称': '上海三菱电梯有限公司',
    '制造商地址': '上海市闵行区紫竹科学园区江川路811号',
    '投标人的单位地址': '深圳市南山区科技园南区高新南一道001号华升大厦18层',
    '国家／地区名称': '中国',
    '设备名称': '乘客电梯 HLK-1000/1.75',
    '标段编号': 'T-2025-001',

    # 法定代表人信息
    '法定代表人姓名': '张伟',
    '法定代表人性别': '男',
    '法定代表人年龄': '48',
    '法定代表人职务': '董事长、总经理',

    # 被委托人信息
    '被委托人姓名': '李明',
    '被委托人性别': '男',
    '被委托人年龄': '42',
    '被委托人职务': '销售总监',
    '被委托人电话': '0755-88886667',
    '被委托人身份证号': '440301198301011234',

    # 项目负责人信息
    '项目负责人姓名': '王建国',
    '项目负责人年龄': '45',
    '项目负责人性别': '男',
    '项目负责人职务': '项目总监',
    '项目负责人职称': '高级工程师',
    '项目负责人学历': '本科',
    '项目负责人执业资格': '一级建造师（机电工程）',
    '项目负责人资格证书编号': '粤1442008200987654',
    '项目负责人拟担任职务': '项目总监',
    '项目负责人毕业年份': '2008',
    '项目负责人毕业学校': '同济大学',
    '项目负责人毕业专业': '机械工程',
    '项目负责人学制': '4',
    '设备规格型号': 'HLK-1000/1.75',

    # 企业基本信息
    '单位性质': '有限责任公司（自然人投资或控股）',
    '地址': '深圳市南山区科技园南区高新南一道001号华升大厦18层',
    '成立时间': '2010年08月18日',
    '经营期限': '2010年08月18日至2030年08月17日',
    '注册资金': '人民币捌仟万元整',
    '邮政编码': '518057',
    '员工总数': '268人',

    # 联系方式
    '电话': '0755-88886666',
    '传真': '0755-88886688',
    '开户银行': '中国建设银行深圳科技园支行',
    '银行账号': '44201501100052501234',

    # 资质信息
    '资质等级': '特种设备安装改造维修许可证 A级',
    '质量管理体系': 'ISO9001:2015质量管理体系认证',

    # 日期信息
    '投标日期年': '2025',
    '投标日期月': '01',
    '投标日期日': '15',
    '授权日期年': '2025',
    '授权日期月': '01',
    '授权日期日': '10',
    '委托期限开始': '2025年01月10日',
    '委托期限结束': '2025年06月30日',

    # 额外信息
    '信誉最低要求': '无重大违法记录，未被列入失信被执行人名单，近三年无重大质量事故',
    '其他补充说明': '本公司财务状况良好，具备履行合同所需的设备和专业技术能力',
    '标段名称': '华美大厦电梯采购及安装',
    '保证金大写': '人民币伍拾万元整',
    '保证金小写': '500000',
    '成员一名称': '深圳市华升电梯设备有限公司',
    '所有成员单位名称': '深圳市华升电梯设备有限公司',
    '联合体名称': '华升电梯独立投标',
    '某成员单位名称': '深圳市华升电梯设备有限公司',
}


def get_paragraph_text(para):
    """获取段落的完整文本（合并所有runs）"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])


def set_paragraph_text(para, new_text):
    """设置段落的完整文本，保留第一个run的格式"""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)


def replace_in_paragraph(para, old, new):
    """在段落中替换文本"""
    full_text = get_paragraph_text(para)
    if old in full_text:
        new_text = full_text.replace(old, new)
        set_paragraph_text(para, new_text)
        return True
    return False


def set_cell_text(cell, text):
    """设置单元格文本"""
    if cell.paragraphs:
        first_para = cell.paragraphs[0]
        if first_para.runs:
            first_para.runs[0].text = text
            for run in first_para.runs[1:]:
                run.text = ''
        else:
            first_para.text = text
    else:
        cell.text = text


def replace_in_cell(cell, old, new):
    """在单元格中替换文本"""
    for para in cell.paragraphs:
        if replace_in_paragraph(para, old, new):
            return True
    return False


def process_table_type1_leader_info(table, data):
    """
    处理法定代表人信息表（表格1）
    格式：姓 | 名：_____性 | 别：_____
    """
    # 第1行：姓名、性别
    set_cell_text(table.rows[0].cells[1], f'名：{data["法定代表人姓名"]}\t性')
    set_cell_text(table.rows[0].cells[2], f'别：{data["法定代表人性别"]}\t')

    # 第2行：年龄、职务
    set_cell_text(table.rows[1].cells[1], f'龄：{data["法定代表人年龄"]}\t职')
    set_cell_text(table.rows[1].cells[2], f'务：{data["法定代表人职务"]}\t')


def process_table_type4_company_basic(table, data):
    """
    处理投标人基本情况表（表格4）
    """
    # 行0: 投标人名称
    set_cell_text(table.rows[0].cells[1], data['投标人名称'])

    # 行1: 注册资金 | [空白] | 成立时间 | [空白]
    set_cell_text(table.rows[1].cells[1], data['注册资金'])
    set_cell_text(table.rows[1].cells[3], data['成立时间'])

    # 行2: 注册地址
    set_cell_text(table.rows[2].cells[1], data['地址'])

    # 行3: 邮政编码、员工总数
    set_cell_text(table.rows[3].cells[1], data['邮政编码'])
    set_cell_text(table.rows[3].cells[3], data['员工总数'])

    # 行4: 联系方式
    set_cell_text(table.rows[4].cells[1], data['电话'])
    set_cell_text(table.rows[4].cells[3], data['传真'])

    # 行5: 开户银行
    set_cell_text(table.rows[5].cells[1], data['开户银行'])

    # 行6: 银行账号
    set_cell_text(table.rows[6].cells[1], data['银行账号'])


def process_table_type5_project_experience(table, data):
    """
    处理近年类似项目情况表（表格5）
    """
    # 项目1的数据
    project_data = {
        '设备名称': data['设备名称'],
        '规格型号': 'HLK-1000/1.75',
        '数量': '6台',
        '合同价格': '人民币叁佰陆拾万元整',
        '项目业主': '深圳科技园物业管理有限公司',
        '合同签订时间': '2023年06月',
        '备注': '已验收合格'
    }

    set_cell_text(table.rows[0].cells[1], project_data['设备名称'])
    set_cell_text(table.rows[1].cells[1], project_data['规格型号'])
    set_cell_text(table.rows[2].cells[1], project_data['数量'])
    set_cell_text(table.rows[3].cells[1], project_data['合同价格'])
    set_cell_text(table.rows[4].cells[1], project_data['项目业主'])
    set_cell_text(table.rows[5].cells[1], project_data['合同签订时间'])
    set_cell_text(table.rows[6].cells[1], project_data['备注'])


def process_table_type6_project_manager(table, data):
    """
    处理项目负责人简历表（表格6）
    """
    # 行0: 姓名、年龄
    set_cell_text(table.rows[0].cells[1], data['项目负责人姓名'])
    set_cell_text(table.rows[0].cells[3], data['项目负责人年龄'])

    # 行1: 职务、职称
    set_cell_text(table.rows[1].cells[1], data['项目负责人职务'])
    set_cell_text(table.rows[1].cells[3], data['项目负责人职称'])

    # 行2: 学历
    set_cell_text(table.rows[2].cells[1], data['项目负责人学历'])


def analyze_doc_structure(doc):
    """分析文档结构，用于调试"""
    print("=" * 60)
    print("文档结构分析")
    print("=" * 60)
    print(f"段落数量: {len(doc.paragraphs)}")
    print(f"表格数量: {len(doc.tables)}")
    print()

    # 分析表格
    for i, table in enumerate(doc.tables):
        print(f"表格{i}: {len(table.rows)}行 x {len(table.rows[0].cells)}列")
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            first_cell_text = get_paragraph_text(table.rows[0].cells[0].paragraphs[0])
            print(f"  首单元格: {first_cell_text[:50]}")
    print()


def get_cell_text(cell):
    """获取单元格文本"""
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""


def fill_table_type4_project_experience(table, data):
    """
    处理近年类似项目情况表（表格4）
    修正：正确处理表格结构
    """
    # 第2行: 规格和型号
    if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
        set_cell_text(table.rows[1].cells[1], data.get('设备规格型号', 'HLK-1000/1.75'))

    # 第2行: 项目名称 (Row2)
    if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
        set_cell_text(table.rows[2].cells[1], '深圳湾科技生态园电梯采购项目')

    # 第3行: 买方名称
    if len(table.rows) > 3 and len(table.rows[3].cells) > 1:
        set_cell_text(table.rows[3].cells[1], '深圳湾科技发展有限公司')

    # 第4行: 买方联系人及电话
    if len(table.rows) > 4 and len(table.rows[4].cells) > 1:
        set_cell_text(table.rows[4].cells[1], '李经理 0755-86543210')

    # 第5行: 合同价格
    if len(table.rows) > 5 and len(table.rows[5].cells) > 1:
        set_cell_text(table.rows[5].cells[1], '人民币叁佰陆拾万元整（¥3,600,000.00）')

    # 第6行: 合同项目负责人
    if len(table.rows) > 6 and len(table.rows[6].cells) > 1:
        set_cell_text(table.rows[6].cells[1], data['项目负责人姓名'])


def process_table_type6_project_manager(table, data):
    """
    处理项目负责人简历表（表格6）- 修正版本
    """
    # 行0: 姓名、年龄、学历
    if len(table.rows) > 0:
        if len(table.rows[0].cells) > 1:
            set_cell_text(table.rows[0].cells[1], data['项目负责人姓名'])
        # 年龄可能在第2或第4列
        for ci in [2, 3, 4]:
            if len(table.rows[0].cells) > ci:
                cell_text = get_cell_text(table.rows[0].cells[ci])
                if '年龄' in cell_text or cell_text.strip() == '' or '年' in cell_text:
                    set_cell_text(table.rows[0].cells[ci], data['项目负责人年龄'])
                    break
        # 学历
        for ci in [5, 6, 7]:
            if len(table.rows[0].cells) > ci:
                cell_text = get_cell_text(table.rows[0].cells[ci])
                if '学历' in cell_text or cell_text.strip() == '':
                    set_cell_text(table.rows[0].cells[ci], data['项目负责人学历'])
                    break

    # 行1: 职称、单位职务、拟在本标段项目担任职务
    if len(table.rows) > 1:
        if len(table.rows[1].cells) > 1:
            set_cell_text(table.rows[1].cells[1], data['项目负责人职称'])
        # 单位职务
        for ci in [2, 3, 4]:
            if len(table.rows[1].cells) > ci:
                cell_text = get_cell_text(table.rows[1].cells[ci])
                if '职务' in cell_text or cell_text.strip() == '':
                    set_cell_text(table.rows[1].cells[ci], data['项目负责人职务'])
                    break
        # 拟担任职务
        for ci in [5, 6, 7]:
            if len(table.rows[1].cells) > ci:
                cell_text = get_cell_text(table.rows[1].cells[ci])
                if '担任职务' in cell_text or '拟在' in cell_text or cell_text.strip() == '':
                    set_cell_text(table.rows[1].cells[ci], data.get('项目负责人拟担任职务', '项目总监'))
                    break

    # 行2: 执业资格、资格证书编号
    if len(table.rows) > 2:
        if len(table.rows[2].cells) > 1:
            set_cell_text(table.rows[2].cells[1], data.get('项目负责人执业资格', '一级建造师（机电工程）'))
        # 资格证书编号
        for ci in [2, 3, 4]:
            if len(table.rows[2].cells) > ci:
                cell_text = get_cell_text(table.rows[2].cells[ci])
                if '资格证书编号' in cell_text or '编号' in cell_text or cell_text.strip() == '':
                    set_cell_text(table.rows[2].cells[ci], data.get('项目负责人资格证书编号', '粤1442008200987654'))
                    break

    # 行3: 毕业学校信息 - 替换整个文本
    if len(table.rows) > 3 and len(table.rows[3].cells) > 1:
        grad_year = data.get('项目负责人毕业年份', '2008')
        grad_school = data.get('项目负责人毕业学校', '同济大学')
        grad_major = data.get('项目负责人毕业专业', '机械工程')
        study_years = data.get('项目负责人学制', '4')
        grad_text = f"\t{grad_year}年  月毕业于\t{grad_school}\t学校\t{grad_major}\t专业，学制\t{study_years}\t年"
        set_cell_text(table.rows[3].cells[1], grad_text)


def fill_special_dates(doc, data):
    """处理特殊日期格式"""
    date_patterns = [
        # 模式: 年 月 日 (法定代表人身份证明末尾)
        ('年 月 日', f"{data['授权日期年']}年 {data['授权日期月']}月 {data['授权日期日']}日"),
        (' 年 月 日', f" {data['授权日期年']}年 {data['授权日期月']}月 {data['授权日期日']}日"),
        ('\t年\t月\t日', f"\t{data['授权日期年']}年\t{data['授权日期月']}月\t{data['授权日期日']}日"),
    ]

    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        new_text = text
        for old, new in date_patterns:
            if old in new_text and '成立时间' not in new_text and '开立时间' not in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(para, new_text)


def fill_tender_document(input_path, output_path, data):
    """填充招标文件"""
    import sys
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

    print(f"正在读取文档: {input_path}")
    doc = Document(input_path)

    # 分析文档结构
    analyze_doc_structure(doc)

    # ============ 第1步：全局括号型占位符替换 ============
    print("步骤1: 替换括号型占位符...")

    bracket_placeholders = [
        ('（招标人名称）', data['招标人名称']),
        ('（招标项目名称）', data['招标项目名称']),
        ('（投标人名称）', data['投标人名称']),
        ('（制造商名称）', data['制造商名称']),
        ('（制造商地址）', data['制造商地址']),
        ('（投标人的单位地址）', data['投标人的单位地址']),
        ('（国家／地区名称）', data['国家／地区名称']),
        ('（设备名称）', data['设备名称']),
        ('（标段编号）', data['标段编号']),
        ('（姓名）', data['法定代表人姓名']),
        ('（盖单位章）', '（盖章）'),
        ('（签字或盖章）', '（签字）'),
        ('（单位负责人）', data['法定代表人姓名']),
        ('（不含报价）', ''),
        ('（信誉最低要求）', data.get('信誉最低要求', '无不良记录')),
        ('（其他补充说明）', data.get('其他补充说明', '无')),
        ('（本项目不适用）', '本项目为独立投标，不适用联合体'),
        ('（或基本存款账户编号）', ''),
        ('（标段名称）', data.get('标段名称', data['招标项目名称'])),
        ('（大写）', data.get('保证金大写', '人民币伍拾万元整')),
        ('（负责人）', data['法定代表人姓名']),
        ('（公章）', ''),
        ('（或授权代表）', data.get('被委托人姓名', data['法定代表人姓名'])),
        ('（签字）', ''),
        ('（或担保机构担保或保证保险）', ''),
        ('（成员一）', data.get('成员一名称', '')),
        ('（所有成员单位名称）', data.get('所有成员单位名称', data['投标人名称'])),
        ('（联合体名称）', data.get('联合体名称', '')),
        ('（某成员单位名称）', data.get('某成员单位名称', data['投标人名称'])),
    ]

    # 在段落中替换
    for para in doc.paragraphs:
        for old, new in bracket_placeholders:
            replace_in_paragraph(para, old, new)

    # 在表格中替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in bracket_placeholders:
                    replace_in_cell(cell, old, new)

    # ============ 第2步：下划线日期替换 ============
    print("步骤2: 替换下划线日期...")

    date_placeholders = [
        ('_年_月_日', f"{data['投标日期年']}年{data['投标日期月']}月{data['投标日期日']}日"),
        ('_年 _月 _日', f"{data['投标日期年']}年 {data['投标日期月']}月 {data['投标日期日']}日"),
        ('_年', f"{data['投标日期年']}年"),
        ('_月', f"{data['投标日期月']}月"),
        ('_日', f"{data['投标日期日']}日"),
    ]

    for para in doc.paragraphs:
        for old, new in date_placeholders:
            replace_in_paragraph(para, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in date_placeholders:
                    replace_in_cell(cell, old, new)

    # ============ 第3步：空白填空型处理 ============
    print("步骤3: 处理空白填空...")

    blank_placeholders = [
        ('投 标 人：', f"投 标 人：{data['投标人名称']}"),
        ('单位性质：', f"单位性质：{data['单位性质']}"),
        ('地    址：', f"地    址：{data['地址']}"),
        ('单位地址：', f"单位地址：{data['地址']}"),
        ('成立时间：', f"成立时间：{data['成立时间']}"),
        ('经营期限：', f"经营期限：{data['经营期限']}"),
        ('邮政编码：', f"邮政编码：{data['邮政编码']}"),
    ]

    for para in doc.paragraphs:
        text = get_paragraph_text(para).strip()
        for prefix, full_text in blank_placeholders:
            if text == prefix or text.startswith(prefix) and len(text) <= len(prefix) + 2:
                set_paragraph_text(para, full_text)
                break

    # ============ 第4步：表格特殊处理 ============
    print("步骤4: 处理表格数据...")

    if len(doc.tables) > 0:
        # 表格1: 法定代表人信息表
        try:
            process_table_type1_leader_info(doc.tables[0], data)
            print("  - 表格1（法定代表人信息表）已填充")
        except Exception as e:
            print(f"  - 表格1处理失败: {e}")

    if len(doc.tables) > 3:
        # 表格4: 投标人基本情况表
        try:
            process_table_type4_company_basic(doc.tables[3], data)
            print("  - 表格4（投标人基本情况表）已填充")
        except Exception as e:
            print(f"  - 表格4处理失败: {e}")

    if len(doc.tables) > 4:
        # 表格4: 近年类似项目情况表
        try:
            fill_table_type4_project_experience(doc.tables[4], data)
            print("  - 表格4（近年类似项目情况表）已填充")
        except Exception as e:
            print(f"  - 表格4处理失败: {e}")

    if len(doc.tables) > 5:
        # 表格5: 项目负责人简历表
        try:
            process_table_type6_project_manager(doc.tables[5], data)
            print("  - 表格5（项目负责人简历表）已填充")
        except Exception as e:
            print(f"  - 表格5处理失败: {e}")

    # ============ 第5步：特殊日期处理 ============
    print("步骤5: 处理特殊日期格式...")
    fill_special_dates(doc, data)

    # ============ 第5步：特殊文本处理 ============
    print("步骤5: 处理特殊文本...")

    # 处理授权委托书（第二个姓名应为被委托人）
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        # 如果在授权委托书部分，且包含"授权委托"相关的姓名
        if '授权委托' in text or '代理人' in text or '被授权人' in text:
            if data['法定代表人姓名'] in text and '委托' in text:
                replace_in_paragraph(para, data['法定代表人姓名'], data['被委托人姓名'])

    # 处理"我"字替换（法定代表人证明书）
    # 这部分通常在"系"字后面的姓名替换
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        # 检查是否是"系"开头的说明
        if '系' in text and data['法定代表人姓名'] in text:
            # 保持原样，因为这是正确的
            pass

    # ============ 保存文档 ============
    print(f"正在保存文档: {output_path}")
    doc.save(output_path)
    print("✓ 文档填充完成！")

    return doc


def main():
    input_file = r"D:\桌面\招标文件正文.docx"
    output_file = r"D:\桌面\招标文件正文_已填充_基于python-docx.docx"

    try:
        fill_tender_document(input_file, output_file, COMPANY_DATA)
        print(f"\n✓ 成功！输出文件: {output_file}")
    except Exception as e:
        print(f"\n✗ 错误: {e}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()
