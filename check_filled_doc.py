#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查填充后的文档
"""

from docx import Document

def get_paragraph_text(para):
    """获取段落的完整文本"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def check_document(path):
    print(f"检查文档: {path}")
    print("=" * 60)

    doc = Document(path)

    # 检查段落中的关键内容
    print("\n[段落内容检查 - 包含填充关键词的段落]")
    keywords = ['深圳市华升', '华美建设', '张伟', '李明', '王建国', '上海三菱']
    found_count = 0

    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para).strip()
        if any(kw in text for kw in keywords) or ('投标' in text and len(text) < 100):
            print(f"  段落{i}: {text[:80]}")
            found_count += 1

    print(f"\n找到 {found_count} 个相关段落")

    # 检查表格内容
    print("\n[表格内容检查]")
    for ti, table in enumerate(doc.tables):
        print(f"\n表格 {ti}: {len(table.rows)}行 x {len(table.rows[0].cells)}列")

        # 显示表格前3行的内容
        for ri, row in enumerate(table.rows[:3]):
            row_text = []
            for ci, cell in enumerate(row.cells[:4]):  # 只显示前4列
                text = get_paragraph_text(cell.paragraphs[0])[:20] if cell.paragraphs else ''
                row_text.append(f"[{text}]")
            print(f"  行{ri}: {' | '.join(row_text)}")

        if len(table.rows) > 3:
            print(f"  ... ({len(table.rows) - 3} 行省略)")

    # 检查是否还有占位符残留
    print("\n[占位符残留检查]")
    placeholders = ['（招标人名称）', '（招标项目名称）', '（投标人名称）',
                    '（制造商名称）', '（姓名）', '_年_月_日']

    for p in placeholders:
        found = False
        for para in doc.paragraphs:
            if p in get_paragraph_text(para):
                found = True
                break
        if not found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if p in get_paragraph_text(cell.paragraphs[0]):
                            found = True
                            break
                    if found:
                        break
                if found:
                    break

        status = "FOUND" if found else "CLEARED"
        print(f"  {p}: {status}")

if __name__ == '__main__':
    check_document(r"D:\桌面\招标文件正文_已填充_基于python-docx.docx")
