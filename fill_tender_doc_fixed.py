#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件自动填充脚本 - 修正版
正确处理合并单元格和表格结构
"""

from docx import Document

def get_paragraph_text(para):
    """获取段落的完整文本"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def set_paragraph_text(para, new_text):
    """设置段落的完整文本"""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)

def set_cell_text(cell, text):
    """设置单元格文本"""
    if cell.paragraphs:
        first_para = cell.paragraphs[0]
        if first_para.runs:
            first_para.runs[0].text = text
            for run in first_para.runs[1:]:
                run.text = ''
        else:
            first_para.text = text
    else:
        cell.text = text

def replace_in_paragraph(para, old, new):
    """在段落中替换文本"""
    full_text = get_paragraph_text(para)
    if old in full_text:
        new_text = full_text.replace(old, new)
        set_paragraph_text(para, new_text)
        return True
    return False

def replace_in_cell(cell, old, new):
    """在单元格中替换文本"""
    for para in cell.paragraphs:
        if replace_in_paragraph(para, old, new):
            return True
    return False

def get_cell_text(cell):
    """获取单元格文本"""
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""

# ============ 虚拟企业投标数据 ============
COMPANY_DATA = {
    '投标人名称': '深圳市华升电梯设备有限公司',
    '招标人名称': '华美建设集团有限公司',
    '招标项目名称': '华美大厦电梯采购及安装项目',
    '制造商名称': '上海三菱电梯有限公司',
    '制造商地址': '上海市闵行区紫竹科学园区江川路811号',
    '投标人的单位地址': '深圳市南山区科技园南区高新南一道001号华升大厦18层',
    '国家／地区名称': '中国',
    '设备名称': '乘客电梯 HLK-1000/1.75',
    '标段编号': 'T-2025-001',

    '法定代表人姓名': '张伟',
    '法定代表人性别': '男',
    '法定代表人年龄': '48',
    '法定代表人职务': '董事长、总经理',

    '被委托人姓名': '李明',
    '被委托人性别': '男',
    '被委托人年龄': '42',
    '被委托人职务': '销售总监',

    '项目负责人姓名': '王建国',
    '项目负责人年龄': '45',
    '项目负责人职称': '高级工程师',

    '单位性质': '有限责任公司（自然人投资或控股）',
    '地址': '深圳市南山区科技园南区高新南一道001号华升大厦18层',
    '成立时间': '2010年08月18日',
    '经营期限': '2010年08月18日至2030年08月17日',
    '注册资金': '人民币捌仟万元整',
    '邮政编码': '518057',
    '员工总数': '268人',

    '联系人': '李明',
    '电话': '0755-88886666',
    '传真': '0755-88886688',
    '开户银行': '中国建设银行深圳科技园支行',
    '银行账号': '44201501100052501234',

    '投标日期年': '2025',
    '投标日期月': '01',
    '投标日期日': '15',
}

def fill_cell_after_label(row, label, value):
    """
    在行中查找标签，并填充标签后的第一个空白单元格
    """
    for ci, cell in enumerate(row.cells):
        text = get_cell_text(cell).strip()
        if label in text and len(text) < len(label) + 5:  # 确保是标签单元格
            # 找到下一个单元格
            if ci + 1 < len(row.cells):
                target_cell = row.cells[ci + 1]
                target_text = get_cell_text(target_cell).strip()
                # 只有当目标单元格为空或是占位符时才填充
                if not target_text or target_text in ['', '　', ' ']:
                    set_cell_text(target_cell, value)
                    return True
    return False

def process_table_type1_leader_info(table, data):
    """处理法定代表人信息表（表格1）"""
    # 第1行：姓名、性别
    set_cell_text(table.rows[0].cells[1], f'名：{data["法定代表人姓名"]}\t性')
    set_cell_text(table.rows[0].cells[2], f'别：{data["法定代表人性别"]}\t')
    # 第2行：年龄、职务
    set_cell_text(table.rows[1].cells[1], f'龄：{data["法定代表人年龄"]}\t职')
    set_cell_text(table.rows[1].cells[2], f'务：{data["法定代表人职务"]}\t')

def process_table_type3_company_basic(table, data):
    """
    处理投标人基本情况表（表格3）- 修正版
    正确处理合并单元格结构
    """
    # 行0: 投标人名称 - cells[1]跨4列
    set_cell_text(table.rows[0].cells[1], data['投标人名称'])

    # 行1: 注册资金 | [span=2] | 成立时间 | [空白]
    # cells[1]是注册资金的数据格, cells[4]是成立时间的数据格
    set_cell_text(table.rows[1].cells[1], data['注册资金'])
    set_cell_text(table.rows[1].cells[4], data['成立时间'])

    # 行2: 注册地址 - cells[1]跨4列
    set_cell_text(table.rows[2].cells[1], data['地址'])

    # 行3: 邮政编码 | [span=2] | 员工总数 | [空白]
    set_cell_text(table.rows[3].cells[1], data['邮政编码'])
    set_cell_text(table.rows[3].cells[4], data['员工总数'])

    # 行4: (空) | 联系人 | [空白] | 电话 | [空白]
    set_cell_text(table.rows[4].cells[2], data['联系人'])
    set_cell_text(table.rows[4].cells[4], data['电话'])

    # 行5: (空) | 网址 | [空白] | 传真 | [空白]
    # 网址留空或不填
    set_cell_text(table.rows[5].cells[4], data['传真'])

    # 行6: 法定代表人 | 姓名 | [空白] | 电话 | [空白]
    set_cell_text(table.rows[6].cells[2], data['法定代表人姓名'])
    set_cell_text(table.rows[6].cells[4], data['电话'])

    # 行8: 基本账户开户银行 - cells[1]跨4列
    set_cell_text(table.rows[8].cells[1], data['开户银行'])

    # 行9: 基本账户银行账号 - cells[1]跨4列
    set_cell_text(table.rows[9].cells[1], data['银行账号'])

def process_table_type4_project_experience(table, data):
    """处理近年类似项目情况表（表格4）"""
    project_data = {
        '设备名称': data['设备名称'],
        '规格型号': 'HLK-1000/1.75',
        '数量': '6台',
        '合同价格': '人民币叁佰陆拾万元整',
        '项目业主': '深圳科技园物业管理有限公司',
        '合同签订时间': '2023年06月',
        '备注': '已验收合格'
    }

    set_cell_text(table.rows[0].cells[1], project_data['设备名称'])
    set_cell_text(table.rows[1].cells[1], project_data['规格型号'])
    set_cell_text(table.rows[2].cells[1], project_data['数量'])
    set_cell_text(table.rows[3].cells[1], project_data['合同价格'])
    set_cell_text(table.rows[4].cells[1], project_data['项目业主'])
    set_cell_text(table.rows[5].cells[1], project_data['合同签订时间'])
    set_cell_text(table.rows[6].cells[1], project_data['备注'])

def process_table_type5_project_manager(table, data):
    """处理项目负责人简历表（表格5）"""
    set_cell_text(table.rows[0].cells[1], data['项目负责人姓名'])
    set_cell_text(table.rows[0].cells[3], data['项目负责人年龄'])
    set_cell_text(table.rows[1].cells[1], '项目总监')
    set_cell_text(table.rows[1].cells[3], data['项目负责人职称'])

def fill_tender_document(input_path, output_path, data):
    """填充招标文件"""
    print(f"Reading document: {input_path}")
    doc = Document(input_path)

    # ============ 第1步：全局占位符替换 ============
    print("Step 1: Replacing bracket placeholders...")

    bracket_placeholders = [
        ('（招标人名称）', data['招标人名称']),
        ('（招标项目名称）', data['招标项目名称']),
        ('（投标人名称）', data['投标人名称']),
        ('（制造商名称）', data['制造商名称']),
        ('（制造商地址）', data['制造商地址']),
        ('（投标人的单位地址）', data['投标人的单位地址']),
        ('（国家／地区名称）', data['国家／地区名称']),
        ('（设备名称）', data['设备名称']),
        ('（标段编号）', data['标段编号']),
    ]

    for para in doc.paragraphs:
        for old, new in bracket_placeholders:
            replace_in_paragraph(para, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in bracket_placeholders:
                    replace_in_cell(cell, old, new)

    # 处理（姓名）占位符 - 需要根据上下文判断
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        if '（姓名）' in text:
            # 在授权委托书中使用被委托人姓名
            if '授权' in text or '委托' in text or '代理人' in text:
                replace_in_paragraph(para, '（姓名）', data['被委托人姓名'])
            else:
                # 其他情况使用法定代表人姓名
                replace_in_paragraph(para, '（姓名）', data['法定代表人姓名'])

    # ============ 第2步：下划线日期替换 ============
    print("Step 2: Replacing date placeholders...")

    date_replacements = [
        ('_年_月_日', f"{data['投标日期年']}年{data['投标日期月']}月{data['投标日期日']}日"),
        ('_年 _月 _日', f"{data['投标日期年']}年 {data['投标日期月']}月 {data['投标日期日']}日"),
        ('_年', f"{data['投标日期年']}年"),
        ('_月', f"{data['投标日期月']}月"),
        ('_日', f"{data['投标日期日']}日"),
    ]

    for para in doc.paragraphs:
        for old, new in date_replacements:
            replace_in_paragraph(para, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in date_replacements:
                    replace_in_cell(cell, old, new)

    # ============ 第3步：空白填空型处理 ============
    print("Step 3: Processing blank fills...")

    blank_mappings = [
        ('投 标 人：', data['投标人名称']),
        ('单位性质：', data['单位性质']),
        ('地    址：', data['地址']),
        ('单位地址：', data['地址']),
        ('成立时间：', data['成立时间']),
        ('经营期限：', data['经营期限']),
        ('邮政编码：', data['邮政编码']),
    ]

    for para in doc.paragraphs:
        text = get_paragraph_text(para).strip()
        for prefix, value in blank_mappings:
            if text == prefix or (text.startswith(prefix) and len(text) <= len(prefix) + 2):
                set_paragraph_text(para, prefix + value)
                break

    # ============ 第4步：表格处理 ============
    print("Step 4: Processing tables...")

    if len(doc.tables) > 0:
        try:
            process_table_type1_leader_info(doc.tables[0], data)
            print("  - Table 0 (Leader info) filled")
        except Exception as e:
            print(f"  - Table 0 error: {e}")

    if len(doc.tables) > 3:
        try:
            process_table_type3_company_basic(doc.tables[3], data)
            print("  - Table 3 (Company basic) filled")
        except Exception as e:
            print(f"  - Table 3 error: {e}")

    if len(doc.tables) > 4:
        try:
            process_table_type4_project_experience(doc.tables[4], data)
            print("  - Table 4 (Project experience) filled")
        except Exception as e:
            print(f"  - Table 4 error: {e}")

    if len(doc.tables) > 5:
        try:
            process_table_type5_project_manager(doc.tables[5], data)
            print("  - Table 5 (Project manager) filled")
        except Exception as e:
            print(f"  - Table 5 error: {e}")

    # ============ 保存文档 ============
    print(f"Saving document: {output_path}")
    doc.save(output_path)
    print("Document filled successfully!")

    return doc

def main():
    input_file = r"D:\桌面\招标文件正文.docx"
    output_file = r"D:\桌面\招标文件正文_已填充_基于python-docx_v2.docx"

    try:
        fill_tender_document(input_file, output_file, COMPANY_DATA)
        print(f"\nSuccess! Output file: {output_file}")
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
