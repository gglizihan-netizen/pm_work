#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证v2版本的填充结果
"""

from docx import Document

def get_paragraph_text(para):
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def verify_v2():
    doc = Document(r"D:\桌面\招标文件正文_已填充_基于python-docx_v2.docx")

    print("=" * 80)
    print("验证v2版本 - 表格3（投标人基本情况表）")
    print("=" * 80)

    table = doc.tables[3]

    key_checks = [
        (0, 1, '投标人名称', '深圳市华升电梯设备有限公司'),
        (1, 1, '注册资金', '人民币捌仟万元整'),
        (1, 4, '成立时间', '2010年08月18日'),
        (2, 1, '注册地址', '深圳市南山区'),
        (3, 1, '邮政编码', '518057'),
        (3, 4, '员工总数', '268人'),
        (4, 2, '联系人', '李明'),
        (4, 4, '电话', '0755-88886666'),
        (5, 4, '传真', '0755-88886688'),
        (6, 2, '法定代表人姓名', '张伟'),
        (6, 4, '法定代表人电话', '0755-88886666'),
        (8, 1, '开户银行', '中国建设银行'),
        (9, 1, '银行账号', '44201501100052501234'),
    ]

    for row_idx, cell_idx, desc, expected in key_checks:
        try:
            actual = get_paragraph_text(table.rows[row_idx].cells[cell_idx].paragraphs[0])
            status = "OK" if expected in actual else "FAIL"
            print(f"[{status}] Row{row_idx},Col{cell_idx} ({desc}): '{actual}'")
        except Exception as e:
            print(f"[ERR] Row{row_idx},Col{cell_idx} ({desc}): Error - {e}")

if __name__ == '__main__':
    verify_v2()
