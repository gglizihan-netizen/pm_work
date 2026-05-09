#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证填充后的文档
"""
from docx import Document
from docx.oxml.ns import qn

def get_paragraph_text(para):
    """获取段落的完整文本（合并所有runs）"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def get_cell_text(cell):
    """获取单元格文本"""
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""

def verify_document(doc_path, output_report=True):
    """验证文档填充情况"""
    doc = Document(doc_path)

    # 检查结果
    report = {
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables),
        'remaining_placeholders': [],
        'filled_samples': [],
    }

    # 检查剩余的占位符
    import re
    bracket_pattern = re.compile(r'（[^）]+）')

    # 检查段落
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        matches = bracket_pattern.findall(text)
        for m in matches:
            if m not in ['（一）', '（二）', '（三）', '（四）', '（五）']:
                report['remaining_placeholders'].append(f"段落{i}: {m}")

        # 记录一些已填充的内容样本
        if i < 60 and len(text) > 10 and '（' not in text:
            report['filled_samples'].append(f"段落{i}: {text[:80]}")

    # 检查表格
    table_samples = []
    for ti, table in enumerate(doc.tables[:6]):
        table_info = f"\n=== 表格 {ti} ==="
        table_samples.append(table_info)
        for ri in range(min(5, len(table.rows))):
            row_texts = []
            for ci, cell in enumerate(table.rows[ri].cells[:5]):
                txt = get_cell_text(cell)[:30]
                row_texts.append(f"[{ci}]{txt}")
            table_samples.append(f"  Row{ri}: {' | '.join(row_texts)}")

    # 写入报告
    report_path = r"D:\桌面\pm_work\verification_report.txt"
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("文档填充验证报告\n")
        f.write("=" * 80 + "\n\n")

        f.write(f"总段落数: {report['total_paragraphs']}\n")
        f.write(f"总表格数: {report['total_tables']}\n\n")

        f.write("【剩余未填充占位符】\n")
        if report['remaining_placeholders']:
            for item in report['remaining_placeholders'][:20]:
                f.write(f"  - {item}\n")
        else:
            f.write("  未发现明显占位符\n")

        f.write("\n【填充样本（前60段）】\n")
        for item in report['filled_samples'][:30]:
            f.write(f"  {item}\n")

        f.write("\n【表格内容预览】\n")
        for line in table_samples:
            f.write(line + "\n")

    print(f"验证报告已保存到: {report_path}")
    return report

if __name__ == '__main__':
    doc_path = r"D:\桌面\招标文件正文_已填充_基于python-docx.docx"
    verify_document(doc_path)
