#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
详细验证填充后的文档内容
"""

from docx import Document

def get_paragraph_text(para):
    """获取段落的完整文本"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def verify_document(path):
    print("=" * 80)
    print("招标文件填充验证报告")
    print("=" * 80)

    doc = Document(path)

    # 1. 检查关键填充内容
    print("\n[1. 关键企业信息检查]")
    expected_content = {
        "深圳市华升电梯设备有限公司": "投标人名称",
        "华美建设集团有限公司": "招标人名称",
        "华美大厦电梯采购及安装": "项目名称",
        "张伟": "法定代表人姓名",
        "李明": "被委托人姓名",
        "王建国": "项目负责人姓名",
        "上海三菱电梯有限公司": "制造商名称",
        "2025年01月15日": "投标日期",
        "T-2025-001": "标段编号",
    }

    all_text = "\n".join([get_paragraph_text(p) for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + get_paragraph_text(cell.paragraphs[0])

    for content, desc in expected_content.items():
        found = content in all_text
        status = "OK" if found else "MISSING"
        print(f"  [{status}] {desc}: {content}")

    # 2. 检查表格填充
    print("\n[2. 表格填充详情]")

    # 表格0 - 法定代表人信息表
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        print("\n  表格0 - 法定代表人信息表:")
        print(f"    行0-姓名: {get_paragraph_text(t0.rows[0].cells[1].paragraphs[0])}")
        print(f"    行0-性别: {get_paragraph_text(t0.rows[0].cells[2].paragraphs[0])}")
        print(f"    行1-年龄: {get_paragraph_text(t0.rows[1].cells[1].paragraphs[0])}")
        print(f"    行1-职务: {get_paragraph_text(t0.rows[1].cells[2].paragraphs[0])}")

    # 表格3 - 投标人基本情况表
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        print("\n  表格3 - 投标人基本情况表:")
        print(f"    投标人名称: {get_paragraph_text(t3.rows[0].cells[1].paragraphs[0])}")
        print(f"    注册资金: {get_paragraph_text(t3.rows[1].cells[1].paragraphs[0])}")
        print(f"    成立时间: {get_paragraph_text(t3.rows[1].cells[3].paragraphs[0])}")
        print(f"    注册地址: {get_paragraph_text(t3.rows[2].cells[1].paragraphs[0])[:40]}...")

    # 表格4 - 近年类似项目情况表
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        print("\n  表格4 - 近年类似项目情况表:")
        print(f"    设备名称: {get_paragraph_text(t4.rows[0].cells[1].paragraphs[0])}")
        print(f"    规格型号: {get_paragraph_text(t4.rows[1].cells[1].paragraphs[0])}")

    # 表格5 - 项目负责人简历表
    if len(doc.tables) > 5:
        t5 = doc.tables[5]
        print("\n  表格5 - 项目负责人简历表:")
        print(f"    姓名: {get_paragraph_text(t5.rows[0].cells[1].paragraphs[0])}")
        print(f"    年龄: {get_paragraph_text(t5.rows[0].cells[3].paragraphs[0])}")
        print(f"    职务: {get_paragraph_text(t5.rows[1].cells[1].paragraphs[0])}")
        print(f"    职称: {get_paragraph_text(t5.rows[1].cells[3].paragraphs[0])}")

    # 3. 检查下划线是否保留
    print("\n[3. 下划线格式检查]")
    underline_count = 0
    for para in doc.paragraphs:
        if '_' in get_paragraph_text(para):
            underline_count += 1
    print(f"  剩余下划线数量: {underline_count}")

    # 4. 检查占位符残留
    print("\n[4. 占位符残留检查]")
    placeholder_patterns = ['（', '）', '_年', '_月', '_日']
    for p in placeholder_patterns:
        count = all_text.count(p)
        if p in ['（', '）']:
            # 括号可能是正常内容，需要判断
            if count < 20:  # 假设正常内容中的括号较少
                print(f"  括号 '{p}' 出现次数: {count} (正常)")
            else:
                print(f"  括号 '{p}' 出现次数: {count} (可能残留)")
        else:
            if count > 0:
                print(f"  警告: '{p}' 出现 {count} 次，可能有残留")
            else:
                print(f"  OK: '{p}' 已完全清除")

    print("\n" + "=" * 80)
    print("验证完成")
    print("=" * 80)

if __name__ == '__main__':
    verify_document(r"D:\桌面\招标文件正文_已填充_基于python-docx.docx")
