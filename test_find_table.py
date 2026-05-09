#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试表格查找功能
"""

from docx import Document

def get_paragraph_text(para):
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def get_cell_text(cell):
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""

def find_tables_by_header(doc, header_keywords):
    """通过表头关键词查找表格"""
    results = []
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        # 检查第一行或前两行作为表头
        for row_idx in range(min(2, len(table.rows))):
            row_text = " ".join([get_cell_text(cell) for cell in table.rows[row_idx].cells])
            print(f"  Table {i}, Row {row_idx}: {row_text[:60]}")
            for kw in header_keywords:
                if kw in row_text:
                    results.append((i, kw))
                    print(f"    -> 匹配关键词: {kw}")
                    break
    return results

doc = Document(r"D:\桌面\招标文件正文.docx")

print("查找 '投标人基本情况' 相关表格:")
keywords = ['投标人基本情况', '供应商基本情况', '单位基本情况', '企业基本情况']
results = find_tables_by_header(doc, keywords)
print(f"\n结果: {results}")

print("\n\n查找 '法定代表人' 相关表格:")
keywords2 = ['法定代表人', '法人信息']
results2 = find_tables_by_header(doc, keywords2)
print(f"\n结果: {results2}")
