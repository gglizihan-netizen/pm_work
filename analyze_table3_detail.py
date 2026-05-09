#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
详细分析表格3（投标人基本情况表）的结构
"""

from docx import Document
from docx.oxml import parse_xml

def get_paragraph_text(para):
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def get_cell_info(cell):
    """获取单元格的详细信息"""
    tc = cell._tc
    # 获取网格跨度信息
    grid_span = 1
    tcPr = tc.tcPr
    if tcPr is not None:
        gridSpan_elem = tcPr.gridSpan
        if gridSpan_elem is not None:
            grid_span = int(gridSpan_elem.val)

    text = get_paragraph_text(cell.paragraphs[0]) if cell.paragraphs else ""
    return {
        'text': text,
        'grid_span': grid_span,
        'actual_cells': len(tc.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'))
    }

def analyze_table3(doc_path):
    doc = Document(doc_path)

    # 找到投标人基本情况表
    for ti, table in enumerate(doc.tables):
        if table.rows:
            first_cell_text = get_paragraph_text(table.rows[0].cells[0].paragraphs[0])
            if "投标人" in first_cell_text and ("名称" in first_cell_text or "基本" in first_cell_text or "情况" in first_cell_text):
                print(f"=" * 80)
                print(f"表格 {ti} - {first_cell_text}")
                print(f"=" * 80)

                for ri, row in enumerate(table.rows):
                    print(f"\n行 {ri}:")
                    print(f"  单元格数量: {len(row.cells)}")

                    for ci, cell in enumerate(row.cells):
                        text = get_paragraph_text(cell.paragraphs[0]).strip()
                        print(f"    cells[{ci}]: '{text}'")

                # 检查我的填充逻辑会填充到哪里
                print("\n" + "-" * 80)
                print("模拟填充（使用我之前的代码逻辑）:")
                print("-" * 80)
                print(f"rows[0].cells[1] 会被填充为: '投标人名称'")
                print(f"  当前内容: '{get_paragraph_text(table.rows[0].cells[1].paragraphs[0])}'")
                print(f"rows[1].cells[1] 会被填充为: '注册资金'")
                print(f"  当前内容: '{get_paragraph_text(table.rows[1].cells[1].paragraphs[0])}'")
                print(f"rows[1].cells[3] 会被填充为: '成立时间'")
                if len(table.rows[1].cells) > 3:
                    print(f"  当前内容: '{get_paragraph_text(table.rows[1].cells[3].paragraphs[0])}'")
                print(f"rows[2].cells[1] 会被填充为: '注册地址'")
                print(f"  当前内容: '{get_paragraph_text(table.rows[2].cells[1].paragraphs[0])}'")
                print(f"rows[3].cells[1] 会被填充为: '邮政编码'")
                print(f"  当前内容: '{get_paragraph_text(table.rows[3].cells[1].paragraphs[0])}'")
                print(f"rows[3].cells[3] 会被填充为: '员工总数'")
                if len(table.rows[3].cells) > 3:
                    print(f"  当前内容: '{get_paragraph_text(table.rows[3].cells[3].paragraphs[0])}'")

                # 显示正确的标签位置
                print("\n" + "-" * 80)
                print("实际标签位置分析:")
                print("-" * 80)
                labels_to_find = ['注册资金', '成立时间', '邮政编码', '员工总数', '联系人', '电话', '网址', '传真', '法定代表人']
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        text = get_paragraph_text(cell.paragraphs[0]).strip()
                        for label in labels_to_find:
                            if label in text:
                                print(f"  '{label}' 在行{ri}, 列{ci}: '{text}'")

if __name__ == '__main__':
    analyze_table3(r"D:\桌面\招标文件正文.docx")
