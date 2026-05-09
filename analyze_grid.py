#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析表格的网格结构（处理合并单元格）
"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def get_paragraph_text(para):
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def analyze_table_grid(table):
    """分析表格的网格结构，包括合并单元格"""
    print(f"表格: {len(table.rows)}行 x {len(table.rows[0].cells)}单元格(可见)")

    # 获取表格的网格宽度
    tblGrid = table._tbl.tblGrid
    grid_cols = len(tblGrid.findall(qn('w:gridCol')))
    print(f"实际网格列数: {grid_cols}")

    for ri, row in enumerate(table.rows):
        print(f"\n行{ri}:")
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.tcPr

            # 获取gridSpan（跨列数）
            grid_span = 1
            if tcPr is not None:
                gridSpan_elem = tcPr.gridSpan
                if gridSpan_elem is not None:
                    grid_span = int(gridSpan_elem.val)

            # 获取vMerge（跨行信息）
            vmerge = None
            if tcPr is not None:
                vMerge_elem = tcPr.vMerge
                if vMerge_elem is not None:
                    vmerge = vMerge_elem.val

            text = get_paragraph_text(cell.paragraphs[0]).strip()[:20]
            span_info = f"span={grid_span}" if grid_span > 1 else ""
            merge_info = f"vmerge={vmerge}" if vmerge else ""
            print(f"  cells[{ci}]: '{text}' {span_info} {merge_info}")

def find_table3_and_analyze(doc_path):
    doc = Document(doc_path)

    for ti, table in enumerate(doc.tables):
        if table.rows:
            first_cell_text = get_paragraph_text(table.rows[0].cells[0].paragraphs[0])
            if "投标人" in first_cell_text:
                print("=" * 80)
                print(f"表格 {ti}: {first_cell_text}")
                print("=" * 80)
                analyze_table_grid(table)

                # 显示正确的填充位置
                print("\n" + "-" * 80)
                print("正确的填充位置应该是:")
                print("-" * 80)

                # 分析每行应该填充的位置
                for ri, row in enumerate(table.rows):
                    labels_and_targets = []
                    for ci, cell in enumerate(row.cells):
                        text = get_paragraph_text(cell.paragraphs[0]).strip()
                        tc = cell._tc
                        tcPr = tc.tcPr
                        grid_span = 1
                        if tcPr is not None:
                            gridSpan_elem = tcPr.gridSpan
                            if gridSpan_elem is not None:
                                grid_span = int(gridSpan_elem.val)

                        # 根据标签找到应该填充的下一个单元格
                        if text in ['注册资金', '成立时间', '邮政编码', '员工总数', '联系人', '电话', '网址', '传真', '姓名']:
                            # 找到这个标签后的下一个单元格索引
                            next_ci = ci + 1
                            if next_ci < len(row.cells):
                                labels_and_targets.append((text, next_ci))

                    if labels_and_targets:
                        print(f"  行{ri}: {labels_and_targets}")

if __name__ == '__main__':
    find_table3_and_analyze(r"D:\桌面\招标文件正文.docx")
