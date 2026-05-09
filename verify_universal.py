#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证通用版填充结果
"""

from docx import Document

def get_paragraph_text(para):
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def get_cell_text(cell):
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""

def verify():
    doc = Document(r"D:\桌面\招标文件正文_已填充_通用版.docx")

    print("=" * 80)
    print("通用版填充结果验证")
    print("=" * 80)

    # 1. 验证括号型占位符
    print("\n[1. 括号型占位符验证]")
    all_text = "\n".join([get_paragraph_text(p) for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + get_cell_text(cell)

    checks = [
        ('深圳市华升电梯设备有限公司', '投标人名称'),
        ('华美建设集团有限公司', '招标人名称'),
        ('上海三菱电梯有限公司', '制造商名称'),
        ('T-2025-001', '标段编号'),
    ]

    for expected, desc in checks:
        found = expected in all_text
        print(f"  [{'OK' if found else 'FAIL'}] {desc}: {expected if found else 'Not found'}")

    # 2. 验证表格填充
    print("\n[2. 表格填充验证]")

    # 表格0 - 法定代表人信息表
    print("\n  表格0 (法定代表人信息表):")
    t0 = doc.tables[0]
    print(f"    Row0 Col1: {get_cell_text(t0.rows[0].cells[1])}")
    print(f"    Row0 Col2: {get_cell_text(t0.rows[0].cells[2])}")

    # 表格3 - 投标人基本情况表
    print("\n  表格3 (投标人基本情况表):")
    t3 = doc.tables[3]
    checks_table3 = [
        (0, 1, '投标人名称', '深圳市华升'),
        (1, 1, '注册资金', '人民币'),
        (1, 4, '成立时间', '2010'),
        (2, 1, '注册地址', '深圳'),
        (3, 1, '邮政编码', '518057'),
        (3, 4, '员工总数', '268'),
    ]
    for ri, ci, desc, expected in checks_table3:
        actual = get_cell_text(t3.rows[ri].cells[ci])
        found = expected in actual
        print(f"    [{'OK' if found else 'FAIL'}] {desc}: {actual[:30] if found else 'Not found'}")

    # 表格5 - 项目负责人简历表
    print("\n  表格5 (项目负责人简历表):")
    t5 = doc.tables[5]
    print(f"    Row0 Col1 (姓名): {get_cell_text(t5.rows[0].cells[1])}")
    print(f"    Row0 Col3 (年龄): {get_cell_text(t5.rows[0].cells[3])}")

    # 3. 验证日期
    print("\n[3. 日期验证]")
    date_found = '2025年01月15日' in all_text or '2025年 01月 15日' in all_text
    print(f"  [{'OK' if date_found else 'FAIL'}] 投标日期已填充")

    print("\n" + "=" * 80)

if __name__ == '__main__':
    verify()
