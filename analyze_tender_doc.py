#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析招标文件文档结构
"""
from docx import Document
from docx.oxml.ns import qn
import re
import sys
import io

# 设置UTF-8输出
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def get_paragraph_text(para):
    """获取段落的完整文本（合并所有runs）"""
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def get_cell_text(cell):
    """获取单元格文本"""
    if cell.paragraphs:
        return get_paragraph_text(cell.paragraphs[0])
    return ""

def get_grid_span(cell):
    """获取单元格跨越的列数"""
    tc = cell._tc
    tcPr = tc.tcPr
    if tcPr is not None:
        gridSpan_elem = tcPr.gridSpan
        if gridSpan_elem is not None:
            return int(gridSpan_elem.val)
    return 1

def analyze_document_structure(doc_path):
    """
    分析文档结构，生成配置建议
    """
    doc = Document(doc_path)

    report = {
        '段落数': len(doc.paragraphs),
        '表格数': len(doc.tables),
        '发现的括号型占位符': [],
        '发现的下划线日期': [],
        '表格结构': [],
        '所有段落内容': [],
    }

    # 1. 分析段落中的占位符
    bracket_pattern = re.compile(r'（[^）]+）')
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        report['所有段落内容'].append((i, text))
        matches = bracket_pattern.findall(text)
        for m in matches:
            if m not in report['发现的括号型占位符']:
                report['发现的括号型占位符'].append(m)

    # 2. 分析下划线日期
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        if '_年' in text or '_月' in text or '_日' in text or '___' in text or '__' in text:
            report['发现的下划线日期'].append((i, text))

    # 3. 分析表格结构
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue

        table_info = {
            '索引': i,
            '行数': len(table.rows),
            '网格列数': len(table._tbl.tblGrid.findall(qn('w:gridCol'))),
            '内容详情': [],
        }

        # 记录前8行内容
        for ri in range(min(8, len(table.rows))):
            row_content = []
            for ci, cell in enumerate(table.rows[ri].cells):
                text = get_cell_text(cell)
                span = get_grid_span(cell)
                row_content.append((ci, text, span))
            table_info['内容详情'].append((ri, row_content))

        report['表格结构'].append(table_info)

    return report, doc

if __name__ == '__main__':
    doc_path = r"D:\桌面\招标文件正文.docx"
    report, doc = analyze_document_structure(doc_path)

    # 写入文件
    output_path = r"D:\桌面\pm_work\document_analysis_report.txt"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("文档结构分析报告\n")
        f.write("=" * 80 + "\n")
        f.write(f"\n【段落数】: {report['段落数']}\n")
        f.write(f"【表格数】: {report['表格数']}\n")

        f.write("\n【发现的括号型占位符】\n")
        for item in report['发现的括号型占位符']:
            f.write(f"  - {item}\n")

        f.write("\n【发现的下划线/空白位置】\n")
        for i, text in report['发现的下划线日期'][:30]:  # 限制数量
            f.write(f"  段落{i}: {text[:100]}\n")

        f.write("\n【表格结构】\n")
        for table_info in report['表格结构']:
            f.write(f"\n  表格 {table_info['索引']}:\n")
            f.write(f"    行数: {table_info['行数']}\n")
            f.write(f"    列数: {table_info['网格列数']}\n")
            for ri, row_content in table_info['内容详情']:
                row_str = " | ".join([f"[{ci}]{text[:30]}(span={span})" for ci, text, span in row_content])
                f.write(f"    Row{ri}: {row_str}\n")

        f.write("\n\n【所有段落内容】\n")
        for i, text in report['所有段落内容']:
            f.write(f"段落{i}: {text}\n")

    print(f"分析报告已保存到: {output_path}")

    # 同时在控制台输出摘要
    print("\n" + "=" * 80)
    print("文档结构分析摘要")
    print("=" * 80)
    print(f"段落数: {report['段落数']}")
    print(f"表格数: {report['表格数']}")
    print(f"\n发现的括号型占位符 ({len(report['发现的括号型占位符'])}个):")
    for item in report['发现的括号型占位符']:
        try:
            print(f"  - {item}")
        except:
            print(f"  - (编码问题)")
