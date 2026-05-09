#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析原文档的表格结构
"""

from docx import Document

def get_paragraph_text(para):
    if not para.runs:
        return para.text or ''
    return ''.join([run.text for run in para.runs])

def analyze_table_structure(doc_path):
    doc = Document(doc_path)

    print("=" * 80)
    print("原文档表格结构分析")
    print("=" * 80)

    for ti, table in enumerate(doc.tables):
        print(f"\n【表格 {ti}】")
        print(f"总行数: {len(table.rows)}, 总列数: {len(table.rows[0].cells) if table.rows else 0}")
        print("-" * 80)

        for ri, row in enumerate(table.rows):
            cells_info = []
            for ci, cell in enumerate(row.cells):
                text = get_paragraph_text(cell.paragraphs[0]).strip() if cell.paragraphs else ""
                cells_info.append(f"[{ci}]{text[:15]}")
            print(f"  行{ri:2d}: {' | '.join(cells_info)}")

        # 如果是投标人基本情况表，详细分析
        first_cell = get_paragraph_text(table.rows[0].cells[0].paragraphs[0]) if table.rows else ""
        if "投标人" in first_cell and "基本情况" in first_cell:
            print("\n  ★ 这是投标人基本情况表，详细分析:")
            for ri, row in enumerate(table.rows):
                print(f"\n    行{ri}:")
                for ci, cell in enumerate(row.cells):
                    text = get_paragraph_text(cell.paragraphs[0]).strip()
                    print(f"      单元格[{ri},{ci}]: '{text}'")

if __name__ == '__main__':
    analyze_table_structure(r"D:\桌面\招标文件正文.docx")
